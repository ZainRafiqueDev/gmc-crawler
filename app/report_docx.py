"""Basic Markdown -> docx export. Not a general Markdown parser - handles
exactly the subset app/report.py's generate_markdown_report/generate_delta_report
produce (#/##/### headers, "- " bullets, "**bold**" leading labels, plain
paragraphs), which is all this needs per the Phase 1 brief ("Output as
Markdown first; add a docx export step once Markdown output is solid").

Typography: a single sans-serif family (Arial, with Helvetica/system
sans-serif fallbacks baked into the .docx theme) throughout, consistent
heading sizes/weights and generous margins - a plain formal-report look
rather than the docx default Calibri styling.
"""
from __future__ import annotations

import html
import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_BOLD_LEADING_RE = re.compile(r"^\*\*(.+?)\*\*(.*)$")
_SEVERITY_TAG_RE = re.compile(r"^(\[(CRITICAL|HIGH|MEDIUM|LOW)\])(.*)$")
_SEVERITY_FIELD_RE = re.compile(r"^(Severity:\s*)(critical|high|medium|low)(.*)$", re.IGNORECASE)
_TABLE_SEPARATOR_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")
# The one shape app/report.py ever emits an image reference in (annotated
# screenshots follow-up, Part 3) - not general inline Markdown image syntax,
# same scope discipline as the table handling above.
_SCREENSHOT_LINE_RE = re.compile(r"^-\s*\*\*Screenshot:\*\*\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")

_FONT_NAME = "Arial"
_INK = RGBColor(0x1E, 0x1E, 0x2E)
_MUTED = RGBColor(0x55, 0x55, 0x66)
_BRAND = RGBColor(0x4F, 0x46, 0xE5)

_SEVERITY_COLOR = {
    "critical": RGBColor(0xB9, 0x1C, 0x1C),
    "high": RGBColor(0xC2, 0x41, 0x0C),
    "medium": RGBColor(0xA1, 0x62, 0x07),
    "low": RGBColor(0x47, 0x55, 0x69),
}


def _set_all_font_slots(rpr, name: str) -> None:
    # python-docx's run.font.name only sets the "ascii"/"hAnsi" font slots;
    # without also setting eastAsia/cs, some renderers (notably Word on
    # certain locales) fall back to the theme's default (Calibri) for those
    # slots, so headings/body can render in two different fonts side by side.
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _set_east_asian_font(run, name: str) -> None:
    _set_all_font_slots(run._element.get_or_add_rPr(), name)


def _style_run(run, *, size: int | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = _FONT_NAME
    _set_east_asian_font(run, _FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = color if color is not None else _INK


def _setup_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = _FONT_NAME
    _set_all_font_slots(normal.element.get_or_add_rPr(), _FONT_NAME)
    normal.font.size = Pt(11)
    normal.font.color.rgb = _INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_heading_styled(doc: Document, text: str, level: int):
    heading = doc.add_heading(level=level)
    heading.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    heading.paragraph_format.space_after = Pt(8)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(text)
    sizes = {1: 22, 2: 16, 3: 13}
    _style_run(run, size=sizes.get(level, 12), bold=True, color=_BRAND if level == 1 else _INK)
    return heading


def _add_severity_aware_runs(paragraph, text: str) -> None:
    """Colors a leading "[CRITICAL]"/"[HIGH]"/... tag (page-by-page compact
    findings) or a "Severity: <level>" field (full finding detail lines) to
    match the same severity color coding used in the web report and PDF -
    everything else in the line renders as normal body text.
    """
    tag_match = _SEVERITY_TAG_RE.match(text)
    if tag_match:
        run = paragraph.add_run(tag_match.group(1))
        _style_run(run, bold=True, color=_SEVERITY_COLOR[tag_match.group(2).lower()])
        rest = tag_match.group(3)
        if rest:
            _style_run(paragraph.add_run(" "))
            _add_runs(paragraph, rest.lstrip())
        return

    field_match = _SEVERITY_FIELD_RE.match(text)
    if field_match:
        _style_run(paragraph.add_run(field_match.group(1)))
        level = field_match.group(2).lower()
        _style_run(paragraph.add_run(level), bold=True, color=_SEVERITY_COLOR[level])
        if field_match.group(3):
            _add_runs(paragraph, field_match.group(3))
        return

    _add_runs(paragraph, text)


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    # app/report.py's Policy-by-Policy Review matrix is the only markdown
    # table this project generates - handled explicitly rather than as
    # general GFM table support, same scope discipline as the rest of this
    # module (a specific known shape, not a general Markdown parser).
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        p = cell.paragraphs[0]
        _style_run(p.add_run(text), bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            _style_run(cell.paragraphs[0].add_run(text))
    doc.add_paragraph()  # spacing after the table


def _add_screenshot(doc: Document, base_dir: Path | None, alt_text: str, rel_path: str) -> None:
    full_path = (base_dir / rel_path) if base_dir is not None else None
    label_p = doc.add_paragraph()
    _style_run(label_p.add_run("Screenshot"), bold=True)
    if full_path is not None and full_path.is_file():
        try:
            doc.add_picture(str(full_path), width=Inches(5.5))
            return
        except Exception:  # noqa: BLE001 - a corrupt/unreadable image degrades to a text fallback, never breaks the export
            pass
    fallback_p = doc.add_paragraph()
    _style_run(fallback_p.add_run(f"(image not available - expected at {rel_path})"), color=_MUTED)


def markdown_to_docx_bytes(markdown_text: str, base_dir: Path | str | None = None) -> bytes:
    base_dir = Path(base_dir) if base_dir is not None else None
    doc = Document()
    _setup_document(doc)

    # app/security/sanitize.py HTML-escapes scraped-content fields (&, <, >
    # -> &amp;/&lt;/&gt;) so they render correctly if the Markdown is ever
    # interpreted as HTML - but a docx run is plain text, not HTML, so
    # without unescaping first, a finding whose evidence contained a
    # literal "&" would show up here as the literal 5-character string
    # "&amp;" (verified live - a real bug, not hypothetical). Safe to do
    # unconditionally: text with no entities passes through html.unescape
    # unchanged, and this project never intentionally puts a literal
    # "&amp;"-shaped string in a report - it only appears when something
    # upstream already escaped a real "&".
    lines = [html.unescape(raw_line).rstrip() for raw_line in markdown_text.splitlines()]

    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue

        if (
            line.lstrip().startswith("|") and i + 1 < len(lines)
            and _TABLE_SEPARATOR_RE.match(lines[i + 1].strip())
        ):
            header = _split_table_row(line)
            body_rows = []
            j = i + 2
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                body_rows.append(_split_table_row(lines[j]))
                j += 1
            _add_table(doc, header, body_rows)
            i = j
            continue

        if line.startswith("### "):
            _add_heading_styled(doc, line[4:].strip(), level=3)
        elif line.startswith("## "):
            heading = _add_heading_styled(doc, line[3:].strip(), level=2)
            _add_heading_border(heading)
        elif line.startswith("# "):
            _add_heading_styled(doc, line[2:].strip(), level=1)
        elif _SCREENSHOT_LINE_RE.match(line.strip()):
            match = _SCREENSHOT_LINE_RE.match(line.strip())
            _add_screenshot(doc, base_dir, match.group(1), match.group(2))
        elif line.lstrip().startswith("- "):
            text = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            _add_severity_aware_runs(p, text)
        else:
            p = doc.add_paragraph()
            _add_severity_aware_runs(p, line.strip())
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_heading_border(heading) -> None:
    # A thin rule under H2 section headers (Critical issues, Page-by-Page
    # Findings, etc.) - mirrors the border under h2 in the web report
    # (prose-h2:border-b) so the same visual hierarchy carries into Word.
    pPr = heading._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "4", qn("w:color"): "D9D9E3",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_runs(paragraph, text: str) -> None:
    """Handles a single leading **bold** span (how app/report.py formats
    finding titles) - not general inline Markdown parsing.
    """
    match = _BOLD_LEADING_RE.match(text)
    if match:
        _style_run(paragraph.add_run(match.group(1)), bold=True)
        if match.group(2):
            _style_run(paragraph.add_run(match.group(2)))
    else:
        _style_run(paragraph.add_run(text))
