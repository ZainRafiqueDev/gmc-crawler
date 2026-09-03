import io

from docx import Document
from PIL import Image

from app.report_docx import markdown_to_docx_bytes


# --- Annotated screenshots (follow-up round, Part 3) ------------------------

def test_screenshot_line_embeds_a_real_picture_when_the_file_exists(tmp_path):
    (tmp_path / "screenshots").mkdir()
    img_path = tmp_path / "screenshots" / "shop-example-llm_prohibited_content-0-20260101-000000.jpg"
    Image.new("RGB", (400, 300), color="red").save(img_path, format="JPEG")

    markdown = "### Some finding\n- **Screenshot:** ![Annotated screenshot for Some finding](screenshots/shop-example-llm_prohibited_content-0-20260101-000000.jpg)\n"
    docx_bytes = markdown_to_docx_bytes(markdown, base_dir=tmp_path)
    doc = Document(io.BytesIO(docx_bytes))

    assert len(doc.inline_shapes) == 1
    assert any(p.text == "Screenshot" for p in doc.paragraphs)


def test_screenshot_line_falls_back_to_text_when_file_missing(tmp_path):
    markdown = "- **Screenshot:** ![alt](screenshots/does-not-exist.jpg)\n"
    docx_bytes = markdown_to_docx_bytes(markdown, base_dir=tmp_path)
    doc = Document(io.BytesIO(docx_bytes))

    assert len(doc.inline_shapes) == 0
    assert any("image not available" in p.text for p in doc.paragraphs)


def test_screenshot_line_without_base_dir_falls_back_to_text():
    markdown = "- **Screenshot:** ![alt](screenshots/foo.jpg)\n"
    docx_bytes = markdown_to_docx_bytes(markdown)  # no base_dir given
    doc = Document(io.BytesIO(docx_bytes))

    assert len(doc.inline_shapes) == 0
    assert any("image not available" in p.text for p in doc.paragraphs)


def test_converts_headers_and_bullets_to_a_valid_docx():
    markdown = (
        "# GMC Compliance Audit Report\n\n"
        "## Executive Summary\n\n"
        "- Pages crawled: 5\n"
        "- **Missing required page: Privacy policy**\n"
        "  - Severity: critical | Confidence: confirmed\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)
    assert docx_bytes[:2] == b"PK"  # docx is a zip archive

    doc = Document(io.BytesIO(docx_bytes))
    paragraph_texts = [p.text for p in doc.paragraphs]
    assert "GMC Compliance Audit Report" in paragraph_texts
    assert "Executive Summary" in paragraph_texts
    assert any("Pages crawled: 5" in t for t in paragraph_texts)
    assert any("Missing required page: Privacy policy" in t for t in paragraph_texts)


def test_bold_leading_span_is_bold_in_docx():
    markdown = "- **Bold Title** rest of line not bold"
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(io.BytesIO(docx_bytes))
    para = doc.paragraphs[0]
    assert para.runs[0].text == "Bold Title"
    assert para.runs[0].bold is True
    assert "rest of line not bold" in para.runs[1].text
    assert not para.runs[1].bold


def test_empty_lines_are_skipped():
    markdown = "# Title\n\n\nSome text\n"
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.paragraphs) == 2


def test_html_entities_from_sanitize_for_report_are_decoded_not_shown_literally():
    """app/security/sanitize.py HTML-escapes scraped-content fields (for the
    Markdown/HTML rendering path) - a docx run is plain text, not HTML, so
    without unescaping, a finding whose evidence contained a literal "&"
    would show up as the literal 5-character string "&amp;" (a real bug,
    caught live before this test was written).
    """
    markdown = (
        "- **Missing terms &amp; conditions**\n"
        "  - Evidence: Price is $5 &amp; up, and &lt;script&gt; shows literally\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(io.BytesIO(docx_bytes))
    texts = [p.text for p in doc.paragraphs]
    assert any("Missing terms & conditions" in t for t in texts)
    assert any("Price is $5 & up, and <script> shows literally" in t for t in texts)
    assert not any("&amp;" in t or "&lt;" in t for t in texts)


def test_markdown_table_renders_as_a_real_table_not_raw_pipe_text():
    """Regression: app.report's Policy-by-Policy Review matrix is a real
    markdown table, but this exporter used to fall through to plain-
    paragraph rendering for every "|"-prefixed line, showing raw
    "| Policy Area | Status |..." text instead of an actual table."""
    markdown = (
        "## Policy-by-Policy Review\n\n"
        "| Policy Area | Status | Findings | Summary |\n"
        "| --- | --- | --- | --- |\n"
        "| Shipping Policy | Pass | 0 | No issues found. |\n"
        "| Privacy Policy | At Risk | 1 | Missing page |\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(io.BytesIO(docx_bytes))

    assert not any(p.text.strip().startswith("|") for p in doc.paragraphs)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["Policy Area", "Status", "Findings", "Summary"]
    assert [c.text for c in table.rows[1].cells] == ["Shipping Policy", "Pass", "0", "No issues found."]
    assert [c.text for c in table.rows[2].cells] == ["Privacy Policy", "At Risk", "1", "Missing page"]
